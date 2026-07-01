from __future__ import annotations

from pathlib import Path


SECTION_KEYWORDS = {
    "goals": ["育人目标", "教学目标", "课程思政", "习近平文化思想", "三进"],
    "overview": ["案例综述", "设计实施", "理念", "思路", "实施过程"],
    "features": ["特色", "创新", "亮点", "推广", "实践价值"],
    "norms": ["教材", "学分", "学时", "课程名称", "申报"],
    "student_analysis": ["学情", "学生基础", "认知", "学习特点"],
    "content_strategy": ["教学内容", "教学策略", "项目式", "任务式", "情境"],
    "evaluation_reflection": ["评价", "考核", "反思", "改进"],
}


def process_document(path: Path | None, document_type: str) -> dict:
    if path is None or not path.exists():
        return {
            "document_type": document_type,
            "exists": False,
            "file_path": None,
            "file_name": None,
            "extracted_text": "",
            "tables_text": "",
            "section_candidates": {},
            "detected_sections": [],
            "extraction_status": "missing",
            "evidence_sufficiency": "不足",
            "warnings": ["材料缺失，不评价对应线上评审模块。"],
            "error": None,
        }

    suffix = path.suffix.lower()
    if suffix == ".doc":
        return _failure(path, document_type, "unsupported", "旧版.doc文件第一版不解析，请转换为docx后重新上传。")
    try:
        if suffix == ".docx":
            text, tables = _extract_docx(path)
        elif suffix == ".pdf":
            text, tables = _extract_pdf(path), ""
        else:
            return _failure(path, document_type, "unsupported", f"不支持的文档格式：{suffix}")
    except Exception as exc:
        return _failure(path, document_type, "parse_failed", str(exc))

    merged_text = "\n".join(part for part in [text, tables] if part).strip()
    sections = _detect_sections(merged_text)
    warnings: list[str] = []
    if not merged_text:
        warnings.append("文档内容为空或未能提取正文，文档证据不足，需要人工复核。")
    if len(merged_text) < 100:
        warnings.append("文档提取文本较少，可能为空表或模板未填写，需要人工复核。")
    return {
        "document_type": document_type,
        "exists": True,
        "file_path": str(path),
        "file_name": path.name,
        "extracted_text": text,
        "tables_text": tables,
        "section_candidates": sections,
        "detected_sections": list(sections.keys()),
        "extraction_status": "success" if merged_text else "empty",
        "evidence_sufficiency": "中" if merged_text and warnings else ("高" if merged_text else "不足"),
        "warnings": warnings,
        "error": None,
    }


def _extract_docx(path: Path) -> tuple[str, str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少python-docx依赖，请先安装requirements.txt。") from exc
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    return "\n".join(paragraphs), "\n".join(table_rows)


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少pypdf依赖，请先安装requirements.txt。") from exc
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(f"[第{index}页]\n{page.extract_text() or ''}")
    return "\n".join(pages).strip()


def _detect_sections(text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for key, keywords in SECTION_KEYWORDS.items():
        matches = [line for line in lines if any(keyword in line for keyword in keywords)]
        if matches:
            sections[key] = matches[:10]
    return sections


def _failure(path: Path, document_type: str, status: str, error: str) -> dict:
    return {
        "document_type": document_type,
        "exists": True,
        "file_path": str(path),
        "file_name": path.name,
        "extracted_text": "",
        "tables_text": "",
        "section_candidates": {},
        "detected_sections": [],
        "extraction_status": status,
        "evidence_sufficiency": "不足",
        "warnings": ["文档证据不足，需要人工复核。"],
        "error": error,
    }
