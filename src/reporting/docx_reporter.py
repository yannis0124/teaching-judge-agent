from __future__ import annotations

from pathlib import Path
from typing import Any


REPORT_NOTICE = (
    "本报告依据当前已提供材料生成辅助评分建议；材料齐全时支持100分评审；"
    "材料缺失时仅对可评审模块给出建议。最终分数应由人工评委结合完整材料确认。"
)


def write_report(
    output_path: Path,
    evidence_package: dict,
    agent_results: dict | None,
    material_only: bool = False,
    stop_reason: str | None = None,
) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("缺少python-docx依赖，请先安装requirements.txt。") from exc

    document = Document()
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"].font.size = Pt(10.5)

    candidate_id = evidence_package.get("candidate_id", "")
    document.add_heading(f"{candidate_id} 辅助评分建议报告", level=1)

    document.add_heading("一、报告说明", level=2)
    document.add_paragraph(REPORT_NOTICE)

    document.add_heading("二、材料完整性", level=2)
    for line in _material_integrity_lines(evidence_package):
        document.add_paragraph(line)

    if material_only or not agent_results:
        document.add_heading("三、处理结论", level=2)
        document.add_paragraph(stop_reason or "材料不足或模型未运行，未生成正式评分建议。")
        _save(document, output_path)
        return

    final = agent_results.get("final_judgement", {})
    scoring_results = agent_results.get("scoring_results", {})

    document.add_heading("三、总分汇总", level=2)
    _add_score_summary(document, final)

    document.add_heading("四、整体性评价", level=2)
    _add_overall_review(document, agent_results.get("overall_review", {}))

    _add_section_evaluations(document, "五、案例整体设计20分分项评价", scoring_results, "case_design")
    _add_section_evaluations(document, "六、教案20分分项评价", scoring_results, "lesson_plan")
    _add_section_evaluations(document, "七、现场教学展示60分分项评价", scoring_results, "live_teaching")

    document.add_heading("八、证据一致性复核", level=2)
    _add_consistency_review(document, agent_results.get("consistency_review", {}))

    document.add_heading("九、偏差审查", level=2)
    _add_bias_review(document, agent_results.get("bias_review", {}))

    document.add_heading("十、人工复核清单", level=2)
    _add_manual_review_checklist(document, _collect_manual_points(scoring_results, agent_results, evidence_package))

    _save(document, output_path)


def _add_score_summary(document: Any, final: dict) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(["模块", "建议分", "满分/状态"]):
        table.rows[0].cells[idx].text = header
    rows = [
        ("案例整体设计", final.get("case_design_score_20"), "20"),
        ("教案", final.get("lesson_plan_score_20"), "20"),
        ("现场教学展示", final.get("live_teaching_score_60"), "60"),
        ("已评审分数/可评审满分", final.get("available_total_score"), final.get("available_max_score")),
    ]
    for name, score, max_score in rows:
        cells = table.add_row().cells
        cells[0].text = str(name)
        cells[1].text = "" if score is None else str(score)
        cells[2].text = "" if max_score is None else str(max_score)
    if final.get("can_score_total_100"):
        document.add_paragraph(f"100分总分建议：{final.get('full_total_score_100')} / 100")
    else:
        document.add_paragraph("当前材料不足以形成100分总评。")
    if final.get("summary"):
        document.add_paragraph(str(final["summary"]))


def _material_integrity_lines(evidence_package: dict) -> list[str]:
    completeness = evidence_package.get("material_completeness", {})
    ppt = evidence_package.get("ppt_evidence", {})
    visual = evidence_package.get("visual_evidence", {})
    transcription = evidence_package.get("transcription_status", {})
    return [
        f"申报表/案例整体设计：{completeness.get('application_form', 'unknown')}",
        f"教案：{completeness.get('lesson_plan', 'unknown')}",
        f"视频：{completeness.get('video', 'unknown')}",
        f"PPT：{completeness.get('ppt', 'unknown')}",
        f"transcript.srt：{completeness.get('transcript', 'unknown')}；{transcription.get('message', '')}",
        f"关键帧：{'已生成' if visual.get('available') else '不足'}",
        f"PPT截图：{ppt.get('image_status', '未生成PPT页面截图。')}",
        f"可评审满分：{completeness.get('available_max_score', 0)}",
    ]


def _add_overall_review(document: Any, review: dict) -> None:
    dimensions = review.get("dimensions", {}) if isinstance(review, dict) else {}
    document.add_heading("4.1 整体摘要", level=3)
    summary = review.get("overall_summary") if isinstance(review, dict) else None
    document.add_paragraph(str(summary or "未生成整体性评价。请重新运行评审流程以生成overall_review.json。"))

    _add_dimension(
        document,
        "4.2 “三进”融合深度",
        dimensions.get("three_entries_integration", {}),
        [
            ("综合判断", "summary"),
            ("主要优势", "strengths"),
            ("主要不足", "weaknesses"),
            ("关键证据", "evidence"),
            ("优化建议", "improvement_suggestions"),
            ("人工复核点", "manual_review_points"),
        ],
    )
    ai_dimension = dimensions.get("ai_application_effectiveness", {})
    if ai_dimension:
        ai_dimension = {"application_type_line": ai_dimension.get("application_type", "证据不足"), **ai_dimension}
    _add_dimension(
        document,
        "4.3 AI应用有效度",
        ai_dimension,
        [
            ("AI应用类型", "application_type_line"),
            ("综合判断", "summary"),
            ("主要优势", "strengths"),
            ("主要不足", "weaknesses"),
            ("关键证据", "evidence"),
            ("优化建议", "improvement_suggestions"),
            ("人工复核点", "manual_review_points"),
        ],
    )
    _add_dimension(
        document,
        "4.4 申报表-教案-PPT-现场展示一致性",
        dimensions.get("material_consistency", {}),
        [
            ("综合判断", "summary"),
            ("一致之处", "consistent_points"),
            ("断点或不足", "inconsistent_points"),
            ("关键证据", "evidence"),
            ("优化建议", "improvement_suggestions"),
            ("人工复核点", "manual_review_points"),
        ],
    )
    _add_dimension(
        document,
        "4.5 职业教育类型特色",
        dimensions.get("vocational_education_characteristics", {}),
        [
            ("综合判断", "summary"),
            ("主要优势", "strengths"),
            ("主要不足", "weaknesses"),
            ("关键证据", "evidence"),
            ("优化建议", "improvement_suggestions"),
            ("人工复核点", "manual_review_points"),
        ],
    )
    _add_dimension(
        document,
        "4.6 美育与文化表达质量",
        dimensions.get("aesthetic_and_cultural_expression", {}),
        [
            ("综合判断", "summary"),
            ("主要优势", "strengths"),
            ("主要不足", "weaknesses"),
            ("关键证据", "evidence"),
            ("优化建议", "improvement_suggestions"),
            ("人工复核点", "manual_review_points"),
        ],
    )
    _add_dimension(
        document,
        "4.7 教学闭环完整度",
        dimensions.get("teaching_closure", {}),
        [
            ("综合判断", "summary"),
            ("已形成的闭环", "completed_links"),
            ("薄弱环节", "missing_or_weak_links"),
            ("关键证据", "evidence"),
            ("优化建议", "improvement_suggestions"),
            ("人工复核点", "manual_review_points"),
        ],
    )
    _add_dimension(
        document,
        "4.8 证据限制与人工复核重点",
        dimensions.get("evidence_limitations_and_review_focus", {}),
        [
            ("证据限制", "evidence_limitations"),
            ("高优先级复核", "high_priority_review_points"),
            ("中优先级复核", "medium_priority_review_points"),
            ("低优先级复核", "low_priority_review_points"),
        ],
    )


def _add_dimension(document: Any, title: str, data: dict, fields: list[tuple[str, str]]) -> None:
    document.add_heading(title, level=3)
    if not data:
        document.add_paragraph("该维度未生成结构化评价。")
        return
    for label, key in fields:
        value = data.get(key)
        if key == "summary":
            document.add_paragraph(f"{label}：{value or '未生成综合判断。'}")
        else:
            _add_list(document, label, value, max_items=3)


def _add_section_evaluations(document: Any, title: str, scoring_results: dict, section: str) -> None:
    document.add_heading(title, level=2)
    items = [item for item in scoring_results.values() if item.get("section") == section]
    if not items:
        document.add_paragraph("该模块无评分结果。")
        return
    for item in items:
        document.add_heading(str(item.get("indicator_name", "")), level=3)
        document.add_paragraph(
            f"满分：{item.get('max_score')}；建议分：{item.get('suggested_score')}；证据充分性：{item.get('evidence_sufficiency')}"
        )
        _add_list(document, "文档证据摘录", item.get("document_evidence", []), empty_text="本项无独立文档证据摘录。", max_items=3)
        _add_list(document, "视频时间戳证据", item.get("timestamp_evidence", []), empty_text="本项无独立视频时间戳证据。", max_items=3)
        _add_list(document, "PPT页码证据", item.get("ppt_page_evidence", []), empty_text="本项无独立PPT页码证据。", max_items=3)
        _add_list(document, "画面关键帧证据", item.get("keyframe_evidence", []), empty_text="本项无独立关键帧证据。", max_items=3)
        _add_list(document, "主要优势", item.get("strengths", []), empty_text="未提炼出明确优势。", max_items=4)
        _add_list(document, "扣分理由", item.get("deduction_reasons", []), empty_text="未发现明显扣分点。", max_items=4)
        _add_list(document, "人工复核点", item.get("manual_review_points", []), empty_text="本项暂无单独人工复核点。", max_items=4)


def _add_consistency_review(document: Any, review: dict) -> None:
    if review.get("format_error"):
        document.add_paragraph("该Agent输出格式异常，需要人工复核。")
    if review.get("summary"):
        document.add_paragraph(str(review["summary"]))
    issues = _as_list(review.get("consistency_issues"))
    if not issues:
        document.add_paragraph("未发现明确一致性问题。")
        return
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, header in enumerate(["复核点", "结论", "证据", "影响", "是否人工复核"]):
        table.rows[0].cells[idx].text = header
    suggestions = _as_list(review.get("score_adjustment_suggestions"))
    manual_points = "；".join(_as_list(review.get("manual_review_points")))
    for issue in issues[:8]:
        cells = table.add_row().cells
        issue_text = _format_item(issue)
        cells[0].text = _shorten(issue_text, 120)
        cells[1].text = "存在不一致或待确认"
        cells[2].text = _shorten(issue_text, 140)
        cells[3].text = _matching_text(issue_text, suggestions) or "可能影响相关分项评分可信度。"
        cells[4].text = "是" if _has_keyword_match(issue_text, manual_points) else "建议复核"


def _add_bias_review(document: Any, review: dict) -> None:
    if review.get("format_error"):
        document.add_paragraph("该Agent输出格式异常，需要人工复核。")
    if review.get("summary"):
        document.add_paragraph(str(review["summary"]))
    risks = _as_list(review.get("bias_risks"))
    if risks:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for idx, header in enumerate(["偏差类型", "是否存在", "风险说明", "对评分的影响", "人工复核建议"]):
            table.rows[0].cells[idx].text = header
        manual_points = "；".join(_as_list(review.get("manual_review_points")))
        for risk in risks[:8]:
            risk_text = _format_item(risk)
            risk_type, risk_body = _split_label(risk_text)
            cells = table.add_row().cells
            cells[0].text = risk_type
            cells[1].text = "存在风险"
            cells[2].text = _shorten(risk_body, 140)
            cells[3].text = "提示相关分项需谨慎采用高分判断。"
            cells[4].text = _matching_manual_point(risk_text, manual_points) or "结合原始材料复核。"
    else:
        document.add_paragraph("未发现明确偏差风险。")
    _add_list(
        document,
        "结构化分数调整建议",
        _normalise_adjustment_suggestions(review.get("score_adjustment_suggestions")),
        empty_text="未生成明确分数调整建议。",
        max_items=8,
    )


def _add_manual_review_checklist(document: Any, points: list[str]) -> None:
    grouped = _group_manual_points(points)
    labels = [
        ("高优先级", grouped["high"]),
        ("中优先级", grouped["medium"]),
        ("低优先级", grouped["low"]),
    ]
    for label, values in labels:
        document.add_paragraph(f"{label}：")
        if not values:
            document.add_paragraph("暂无。")
            continue
        for index, value in enumerate(values, 1):
            document.add_paragraph(f"{index}. {value}")


def _add_list(
    document: Any,
    title: str,
    values: list | dict | str | None,
    empty_text: str = "未生成相关内容。",
    max_items: int | None = None,
) -> None:
    document.add_paragraph(f"{title}：")
    items = _as_list(values)
    if not items:
        document.add_paragraph(empty_text)
        return
    if max_items is not None:
        items = items[:max_items]
    for value in items:
        document.add_paragraph(_shorten(_format_item(value), 260), style="List Bullet")


def _as_list(value: list | dict | str | None) -> list:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return [{key: item} for key, item in value.items()]
    if isinstance(value, str):
        return [value] if value.strip() else []
    return [value]


def _format_item(value) -> str:
    if isinstance(value, dict):
        return "；".join(f"{key}: {_format_item(item)}" for key, item in value.items())
    if isinstance(value, list):
        return "；".join(_format_item(item) for item in value)
    return str(value)


def _shorten(text: str, limit: int) -> str:
    text = " ".join(str(text).split())
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def _split_label(text: str) -> tuple[str, str]:
    for sep in ["：", ":"]:
        if sep in text:
            left, right = text.split(sep, 1)
            return left.strip(), right.strip()
    return text[:20], text


def _normalise_adjustment_suggestions(value) -> list:
    items = _as_list(value)
    normalised = []
    for item in items:
        if isinstance(item, dict):
            text = _format_item(item)
        else:
            text = str(item)
        if text.strip():
            normalised.append(text)
    return normalised


def _matching_text(issue: str, suggestions: list) -> str:
    for suggestion in suggestions:
        text = _format_item(suggestion)
        if _has_keyword_match(issue, text):
            return _shorten(text, 120)
    return ""


def _matching_manual_point(risk: str, manual_points: str) -> str:
    if _has_keyword_match(risk, manual_points):
        return _shorten(manual_points, 120)
    return ""


def _has_keyword_match(left: str, right: str) -> bool:
    keywords = ["AI", "美育", "小组", "研讨", "职业", "岗位", "评价", "PPT", "思政"]
    return any(keyword in left and keyword in right for keyword in keywords)


def _collect_manual_points(scoring_results: dict, agent_results: dict, evidence_package: dict) -> list[str]:
    points: list[str] = []
    points.extend(evidence_package.get("material_integrity", {}).get("notes", []))
    for item in scoring_results.values():
        points.extend(str(point) for point in item.get("manual_review_points", []))
    for key in ["consistency_review", "bias_review", "final_judgement"]:
        points.extend(str(point) for point in agent_results.get(key, {}).get("manual_review_points", []))
    overall_focus = (
        agent_results.get("overall_review", {})
        .get("dimensions", {})
        .get("evidence_limitations_and_review_focus", {})
    )
    for field in ["high_priority_review_points", "medium_priority_review_points", "low_priority_review_points"]:
        points.extend(str(point) for point in overall_focus.get(field, []))
    return points


def _group_manual_points(points: list[str]) -> dict[str, list[str]]:
    grouped = {"high": [], "medium": [], "low": []}
    seen: set[str] = set()
    for point in points:
        canonical = _canonical_manual_point(point)
        if not canonical:
            continue
        if canonical in seen:
            continue
        seen.add(canonical)
        grouped[_manual_priority(canonical)].append(canonical)
    return grouped


def _canonical_manual_point(point: str) -> str:
    text = str(point).strip()
    if not text:
        return ""
    if any(keyword in text for keyword in ["AI", "人工智能", "情感分析", "史料分析"]):
        return "AI情感分析、AI史料分析是否由学生主动使用，并形成分析、判断、修正、评价过程。"
    if "美育" in text:
        return "美育元素是否在课堂中显性化展开，而不只停留在教案设计或开场表述。"
    if any(keyword in text for keyword in ["小组", "研讨", "讨论", "阶梯式"]):
        return "教案中的小组研讨或阶梯式提问是否在现场真实发生，并达到应有深度。"
    if any(keyword in text for keyword in ["职业", "岗位", "场景", "实训", "双师"]):
        return "职业场景、岗位任务或实训活动是否有更具体的课堂支撑。"
    if any(keyword in text for keyword in ["学情", "数据来源", "86%"]):
        return "学情分析数据来源是否可核验。"
    if any(keyword in text for keyword in ["增值性评价", "教学评价", "成长数据"]):
        return "增值性评价是否有具体实施证据。"
    if any(keyword in text for keyword in ["申报表", "第6页", "格式", "标题", "缺失"]):
        return "个别材料格式、标题或正文缺失问题需核对原始文件。"
    if any(keyword in text for keyword in ["H5", "课后", "网页"]):
        return "课后H5学习网页的实际使用效果需结合补充材料确认。"
    return text


def _manual_priority(point: str) -> str:
    if any(keyword in point for keyword in ["AI", "美育", "小组研讨", "阶梯式提问"]):
        return "high"
    if any(keyword in point for keyword in ["职业", "岗位", "学情", "增值性评价", "教学评价"]):
        return "medium"
    return "low"


def _save(document: Any, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
